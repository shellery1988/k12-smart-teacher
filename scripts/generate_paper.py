#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能老师 - 试卷生成脚本

目标：
1. 至少为九大学科生成可用的基础/提高/挑战三层练习。
2. 输入或依赖异常时给中文可操作提示，避免直接抛出英文堆栈。
3. 按 --output 真实写入文件；优先生成 docx，依赖缺失时降级为 Markdown。
"""

import argparse
import json
import sys
from pathlib import Path


SUPPORTED_SUBJECTS = ["语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "政治"]
LEVELS = [
    ("basic", "基础巩固", "⭐"),
    ("improve", "能力提高", "⭐⭐"),
    ("challenge", "拓展挑战", "⭐⭐⭐"),
]


class PaperGenerationError(Exception):
    """面向用户的试卷生成错误。"""


class FriendlyArgumentParser(argparse.ArgumentParser):
    def error(self, message):
        print(f"参数错误：{message}", file=sys.stderr)
        print("请检查：学科、知识点、学生、年级、输出路径是否都已填写。", file=sys.stderr)
        print("示例：python3 scripts/generate_paper.py --subject 数学 --topic 最大公因数 --student 小明 --grade 五年级 --output 练习.md", file=sys.stderr)
        raise SystemExit(2)


def normalize_text(value, default):
    value = (value or "").strip()
    return value if value else default


def detect_stage(grade):
    if any(word in grade for word in ["初", "七", "八", "九"]):
        return "初中"
    if any(word in grade for word in ["高", "十", "十一", "十二"]):
        return "高中"
    if any(word in grade for word in ["一", "二", "三", "四", "五", "六", "小学"]):
        return "小学"
    return "通用"


def q(qtype, content, answer="", hint=""):
    item = {"type": qtype, "content": content}
    if answer:
        item["answer"] = answer
    if hint:
        item["hint"] = hint
    return item


def trim_questions(questions, basic_count, improve_count, challenge_count):
    limits = {
        "basic": max(0, basic_count),
        "improve": max(0, improve_count),
        "challenge": max(0, challenge_count),
    }
    return {level: questions.get(level, [])[:limit] for level, limit in limits.items()}


def generate_math_questions(topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    questions = {"basic": [], "improve": [], "challenge": []}

    if "最大公因数" in topic or "公因数" in topic:
        questions["basic"] = [
            q("计算", "用短除法求下列各组数的最大公因数。\n（1）18 和 27　（2）24 和 36　（3）45 和 60", "9；12；15"),
            q("填空", "如果 a 是 b 的倍数，那么 a 和 b 的最大公因数是（　　），最小公倍数是（　　）。", "b；a"),
            q("判断", "判断对错。\n（1）两个数的最大公因数一定比这两个数都小。（　　）\n（2）如果两个数互质，它们的最大公因数是1。（　　）", "×；√"),
        ]
        questions["improve"] = [
            q("应用题", "有两根铁丝，一根长 72 厘米，另一根长 54 厘米。要截成同样长的小段且没有剩余，每小段最长是多少厘米？一共可以截成多少段？", "18厘米；7段", "同样长、没有剩余、最长 -> 最大公因数。"),
            q("应用题", "48 本故事书和 64 本科技书平均分给若干个班，每班两类书数量都相同。最多可以分给多少个班？每班各多少本？", "16个班；故事书3本、科技书4本"),
            q("应用题", "长 84 厘米、宽 60 厘米的长方形彩纸剪成同样大的正方形且没有剩余，正方形边长最大是多少？最少剪成多少个？", "12厘米；35个"),
        ]
        questions["challenge"] = [
            q("思维题", "已知 a 和 b 的最大公因数是 12，最小公倍数是 72。如果 a = 36，那么 b = ？", "24"),
            q("思维题", "一个数除 60 余 4，除 80 也余 4。这个数最大是多少？", "28", "这个数能同时整除 60-4 和 80-4。"),
        ]
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    questions["basic"] = [
        q("概念填空", f"用自己的话解释“{topic}”的含义，并写出一个简单例子。"),
        q("基础计算", f"围绕“{topic}”设计并完成一道基础计算题，写出关键步骤。"),
        q("易错辨析", f"判断：学习“{topic}”时，只要记住公式就一定能做对所有题。（　　）", "×"),
    ]
    questions["improve"] = [
        q("应用题", f"把“{topic}”放进一个购物、路程或图形场景中，列式并解答。"),
        q("方法比较", f"同一道“{topic}”题目尝试用两种方法解决，并说明哪种更简洁。"),
        q("错因分析", f"写出一道“{topic}”相关错题，指出错误原因并改正。"),
    ]
    questions["challenge"] = [
        q("综合题", f"设计一道包含两个条件转换的“{topic}”综合题，并完整解答。"),
        q("开放题", f"给出一个生活中的“{topic}”问题，说明数学模型、解题步骤和答案。"),
    ]
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_chinese_questions(topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    questions = {
        "basic": [
            q("字词积累", f"围绕“{topic}”整理 5 个关键词，分别写出意思或近义词。"),
            q("句子理解", f"用“{topic}”相关内容造 2 个完整句子，要求表达清楚。"),
            q("基础阅读", f"阅读一段与“{topic}”相关的短文，找出中心句并说明理由。"),
        ],
        "improve": [
            q("阅读分析", f"给一段“{topic}”相关材料，概括主要内容，并找出一个细节依据。"),
            q("表达训练", f"围绕“{topic}”写一段 120 字左右的小短文，要求有开头、经过、结尾。"),
            q("错题修正", f"修改一个与“{topic}”有关的病句，并说明病因。"),
        ],
        "challenge": [
            q("综合赏析", f"选择一段与“{topic}”相关的文字，从词语、修辞或情感中任选两点赏析。"),
            q("迁移写作", f"以“{topic}”为核心，写一个提纲：标题、中心、材料、结尾各一项。"),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_english_questions(topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    questions = {
        "basic": [
            q("Vocabulary", f"Write 5 words or phrases about '{topic}' and give one Chinese meaning for each."),
            q("Sentence", f"Make 3 simple sentences about '{topic}'."),
            q("Choice", f"Choose the correct word to complete a sentence about '{topic}', then explain why."),
        ],
        "improve": [
            q("Grammar", f"Write 3 sentences about '{topic}' using the target grammar point, then mark the key structure."),
            q("Reading", f"Read a short paragraph about '{topic}' and answer: Who/What/Why?"),
            q("Correction", f"Find and correct 3 common mistakes in sentences about '{topic}'."),
        ],
        "challenge": [
            q("Writing", f"Write a 60-100 word paragraph about '{topic}' with at least 3 linking words."),
            q("Speaking", f"Prepare a 1-minute oral answer about '{topic}', including opinion and reason."),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_science_questions(subject, topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    unit = {
        "物理": ("概念/公式", "实验或现象", "单位和条件"),
        "化学": ("概念/方程式", "实验现象", "反应条件和守恒"),
        "生物": ("概念/结构", "生命现象", "条件和变量"),
    }[subject]
    questions = {
        "basic": [
            q("概念解释", f"解释“{topic}”中的一个核心{unit[0]}，并举一个例子。"),
            q("基础判断", f"判断一个关于“{topic}”的说法是否正确，并写出依据。"),
            q("图表整理", f"画出或列出“{topic}”的关键结构/步骤/关系图。"),
        ],
        "improve": [
            q("现象分析", f"描述一个与“{topic}”有关的{unit[1]}，说明原因。"),
            q("实验设计", f"设计一个验证“{topic}”的小实验，写出器材、步骤和观察点。"),
            q("易错辨析", f"列出学习“{topic}”时最容易混淆的两个概念，并比较不同点。"),
        ],
        "challenge": [
            q("综合应用", f"给出一个真实情境，运用“{topic}”解释结果，并标明{unit[2]}。"),
            q("探究题", f"围绕“{topic}”提出一个可探究问题，设计变量控制方案。"),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_humanities_questions(subject, topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    focus = {
        "历史": ("时间、人物、事件", "原因和影响", "史料"),
        "地理": ("位置、气候、地形", "成因和影响", "地图或数据"),
        "政治": ("概念、观点、材料", "原因和做法", "时事材料"),
    }[subject]
    questions = {
        "basic": [
            q("基础梳理", f"整理“{topic}”中的 3 个关键词，并解释含义。"),
            q("信息提取", f"从一段关于“{topic}”的材料中提取{focus[0]}。"),
            q("判断说明", f"判断一个关于“{topic}”的说法是否正确，并写出理由。"),
        ],
        "improve": [
            q("材料分析", f"阅读一则关于“{topic}”的材料，概括核心观点并找出依据。"),
            q("因果分析", f"分析“{topic}”的{focus[1]}，至少写出两点。"),
            q("比较归纳", f"把“{topic}”与一个相近知识点作比较，列出相同点和不同点。"),
        ],
        "challenge": [
            q("综合探究", f"结合{focus[2]}，说明“{topic}”在现实或考试题中的应用。"),
            q("开放表达", f"围绕“{topic}”写一段 150 字左右的小论述，要求观点明确、依据充分。"),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_questions(subject, topic, grade, basic_count, improve_count, challenge_count):
    if subject == "数学":
        return generate_math_questions(topic, grade, basic_count, improve_count, challenge_count)
    if subject == "语文":
        return generate_chinese_questions(topic, grade, basic_count, improve_count, challenge_count)
    if subject == "英语":
        return generate_english_questions(topic, grade, basic_count, improve_count, challenge_count)
    if subject in ["物理", "化学", "生物"]:
        return generate_science_questions(subject, topic, grade, basic_count, improve_count, challenge_count)
    if subject in ["历史", "地理", "政治"]:
        return generate_humanities_questions(subject, topic, grade, basic_count, improve_count, challenge_count)
    raise PaperGenerationError(f"暂不支持“{subject}”。可用学科：{'、'.join(SUPPORTED_SUBJECTS)}。")


def validate_questions(questions):
    total = sum(len(items) for items in questions.values())
    if total == 0:
        raise PaperGenerationError("没有生成任何题目。请换一个更明确的知识点，例如“最大公因数”“现在完成时”“牛顿第二定律”。")


def build_output_data(args, questions):
    return {
        "subject": args.subject,
        "topic": args.topic,
        "student": args.student,
        "grade": args.grade,
        "stage": detect_stage(args.grade),
        "questions": questions,
        "usage_note": "题目为基础模板生成，正式给孩子使用前建议老师或家长快速浏览一遍。"
    }


def render_markdown(data):
    lines = [
        f"# {data['student']}的{data['subject']}练习",
        "",
        f"- 年级：{data['grade']}",
        f"- 知识点：{data['topic']}",
        f"- 学段：{data['stage']}",
        "",
        "> 先回顾讲解，再完成练习；遇到不会的题，先写出想到的步骤。",
        "",
    ]
    for key, title, star in LEVELS:
        lines.extend([f"## {title} {star}", ""])
        for index, item in enumerate(data["questions"].get(key, []), 1):
            lines.append(f"{index}. 【{item['type']}】{item['content']}")
            if item.get("hint"):
                lines.append(f"   - 提示：{item['hint']}")
            lines.append("")
    lines.extend(["## 参考答案", ""])
    for key, title, _ in LEVELS:
        answer_items = [item for item in data["questions"].get(key, []) if item.get("answer")]
        if not answer_items:
            continue
        lines.append(f"### {title}")
        for index, item in enumerate(answer_items, 1):
            lines.append(f"{index}. {item['answer']}")
        lines.append("")
    lines.append(f"> {data['usage_note']}")
    return "\n".join(lines)


def write_docx(path, data):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise PaperGenerationError(
            "当前环境缺少 python-docx，无法生成 Word 文件。已建议改用 .md 输出，或先运行：python3 scripts/quick_setup.py"
        ) from exc

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    document.add_heading(f"{data['student']}的{data['subject']}练习", 0)
    document.add_paragraph(f"年级：{data['grade']}    学段：{data['stage']}    知识点：{data['topic']}")
    document.add_paragraph("先回顾讲解，再完成练习；遇到不会的题，先写出想到的步骤。")

    for key, title, star in LEVELS:
        document.add_heading(f"{title} {star}", level=1)
        for index, item in enumerate(data["questions"].get(key, []), 1):
            document.add_paragraph(f"{index}. 【{item['type']}】{item['content']}")
            if item.get("hint"):
                document.add_paragraph(f"提示：{item['hint']}")

    document.add_heading("参考答案", level=1)
    has_answers = False
    for key, title, _ in LEVELS:
        answer_items = [item for item in data["questions"].get(key, []) if item.get("answer")]
        if not answer_items:
            continue
        has_answers = True
        document.add_heading(title, level=2)
        for index, item in enumerate(answer_items, 1):
            document.add_paragraph(f"{index}. {item['answer']}")
    if not has_answers:
        document.add_paragraph("开放题暂无唯一答案，请根据讲解要点和表达完整度评价。")

    document.add_paragraph(data["usage_note"])
    document.save(path)


def write_output(output_path, data):
    path = Path(output_path).expanduser()
    path.parent.mkdir(parents=True, exist_ok=True)
    suffix = path.suffix.lower()

    if suffix == ".json":
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return path
    if suffix == ".docx":
        try:
            write_docx(path, data)
            return path
        except PaperGenerationError:
            fallback = path.with_suffix(".md")
            fallback.write_text(render_markdown(data), encoding="utf-8")
            print(f"提示：Word 依赖不可用，已自动生成 Markdown 文件：{fallback}")
            return fallback

    if suffix not in [".md", ".txt"]:
        path = path.with_suffix(".md")
    path.write_text(render_markdown(data), encoding="utf-8")
    return path


def positive_int(value):
    try:
        number = int(value)
    except ValueError as exc:
        raise argparse.ArgumentTypeError("题目数量必须是整数。") from exc
    if number < 0 or number > 10:
        raise argparse.ArgumentTypeError("每层题目数量建议在 0-10 之间。")
    return number


def parse_args(argv):
    parser = FriendlyArgumentParser(
        description="生成K12三层练习题（支持九大学科，输出 docx/md/json）",
        epilog="示例：python3 scripts/generate_paper.py --subject 数学 --topic 最大公因数 --student 小明 --grade 五年级 --output 练习/小明.docx",
    )
    parser.add_argument("--subject", required=True, help=f"学科：{'、'.join(SUPPORTED_SUBJECTS)}")
    parser.add_argument("--topic", required=True, help="知识点主题，例如：最大公因数、阅读理解、现在完成时")
    parser.add_argument("--student", required=True, help="学生姓名")
    parser.add_argument("--grade", required=True, help="年级，例如：小学五年级、初二、高一")
    parser.add_argument("--output", required=True, help="输出文件路径，支持 .docx/.md/.json")
    parser.add_argument("--basic-count", type=positive_int, default=3, help="基础巩固题数量，默认3")
    parser.add_argument("--improve-count", type=positive_int, default=3, help="能力提高题数量，默认3")
    parser.add_argument("--challenge-count", type=positive_int, default=2, help="拓展挑战题数量，默认2")
    args = parser.parse_args(argv)
    if args.subject not in SUPPORTED_SUBJECTS:
        raise PaperGenerationError(f"暂不支持“{args.subject}”。可用学科：{'、'.join(SUPPORTED_SUBJECTS)}。")
    return args


def main(argv=None):
    try:
        args = parse_args(argv or sys.argv[1:])
        args.topic = normalize_text(args.topic, "综合复习")
        args.student = normalize_text(args.student, "同学")
        args.grade = normalize_text(args.grade, "未提供年级")

        questions = generate_questions(
            args.subject,
            args.topic,
            args.grade,
            args.basic_count,
            args.improve_count,
            args.challenge_count,
        )
        validate_questions(questions)
        data = build_output_data(args, questions)
        saved_path = write_output(args.output, data)

        print(json.dumps(data, ensure_ascii=False, indent=2))
        print(f"\n练习已生成：{saved_path}")
        return 0
    except KeyboardInterrupt:
        print("\n已取消生成。", file=sys.stderr)
        return 130
    except PaperGenerationError as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 2
    except Exception as exc:
        print(f"生成失败：遇到未预期问题：{exc}", file=sys.stderr)
        print("建议：检查学科、知识点、输出路径是否正确；如仍失败，请把这条提示发给智能老师。", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
